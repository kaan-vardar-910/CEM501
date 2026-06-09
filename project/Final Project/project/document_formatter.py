"""Format generated CEM document text into styled Word and PDF files.

Turns the plain text produced by the document generators (RFI, submittal,
delay notice, ...) into professional .docx and .pdf files with a consistent
CEM house style: A4 page, running header/footer, a navy title block, bold
field labels, numbered lists and an automatic signature block.

PDF generation tries three strategies in order: LibreOffice (best fidelity),
the ``docx2pdf`` library, then a plain-text ``reportlab`` fallback.

Run ``python document_formatter.py`` to emit a sample .docx and .pdf into
``outputs/``.
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile
from datetime import datetime

try:
    from dotenv import load_dotenv

    load_dotenv()
except ImportError:  # python-dotenv is optional here.
    pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# CEM house colour for titles and field labels.
NAVY = RGBColor(0x1F, 0x38, 0x64)
NAVY_HEX = "1F3864"

# A4 dimensions and the usable text width after 2.5 cm margins on each side.
PAGE_W_MM = 210
PAGE_H_MM = 297
MARGIN_MM = 25
USABLE_W_MM = PAGE_W_MM - (2 * MARGIN_MM)  # 160 mm

OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "outputs"
)


# --------------------------------------------------------------------------- #
# Low-level docx helpers
# --------------------------------------------------------------------------- #
def _add_field(paragraph, field_code: str) -> None:
    """Append a Word field (e.g. PAGE / NUMPAGES) to ``paragraph``."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _horizontal_rule(paragraph) -> None:
    """Draw a thin navy rule along the bottom border of ``paragraph``."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), NAVY_HEX)
    borders.append(bottom)
    p_pr.append(borders)


def _add_tab_stops(paragraph, stops: list[tuple[int, object]]) -> None:
    """Configure tab stops; ``stops`` is a list of (millimetres, alignment)."""
    for position_mm, alignment in stops:
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Mm(position_mm), alignment
        )


def _style_base_font(document: Document) -> None:
    """Set the document-wide Normal font to 11pt Calibri."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _setup_page(document: Document) -> None:
    """Apply A4 size and 2.5 cm margins to every section."""
    for section in document.sections:
        section.page_width = Mm(PAGE_W_MM)
        section.page_height = Mm(PAGE_H_MM)
        section.left_margin = Mm(MARGIN_MM)
        section.right_margin = Mm(MARGIN_MM)
        section.top_margin = Mm(MARGIN_MM)
        section.bottom_margin = Mm(MARGIN_MM)


def _build_header(document: Document, doc_type: str, settings: dict) -> None:
    """Build the running header: project name (left) and doc type + date."""
    today = datetime.now().strftime("%Y-%m-%d")
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    _add_tab_stops(paragraph, [(USABLE_W_MM, WD_TAB_ALIGNMENT.RIGHT)])
    left = settings.get("project_name", "")
    right = f"{doc_type} | {today}"
    run = paragraph.add_run(f"{left}\t{right}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    _horizontal_rule(paragraph)


def _build_footer(document: Document, settings: dict) -> None:
    """Build the running footer: contractor / Page X of Y / contract number."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    _add_tab_stops(
        paragraph,
        [
            (USABLE_W_MM // 2, WD_TAB_ALIGNMENT.CENTER),
            (USABLE_W_MM, WD_TAB_ALIGNMENT.RIGHT),
        ],
    )
    paragraph.add_run(f"{settings.get('contractor_name', '')}\t")
    paragraph.add_run("Page ")
    _add_field(paragraph, "PAGE")
    paragraph.add_run(" of ")
    _add_field(paragraph, "NUMPAGES")
    paragraph.add_run(f"\t{settings.get('contract_number', '')}")
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _build_title_block(
    document: Document, doc_type: str, settings: dict
) -> None:
    """Write the centred company name + document type heading."""
    company = document.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run(settings.get("contractor_name", ""))
    run.bold = True
    run.font.size = Pt(14)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(doc_type)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    _horizontal_rule(document.add_paragraph())


_FIELD_LABEL_RE = re.compile(r"^[^a-z]*:$")
_NUMBERED_RE = re.compile(r"^\s*\d+\.\s+")
_RULE_RE = re.compile(r"^\s*(={3,}|-{3,})\s*$")


def _is_field_label(line: str) -> bool:
    """True if ``line`` looks like an ALL-CAPS or Title-Case field label."""
    stripped = line.strip()
    if not stripped.endswith(":") or len(stripped) > 60:
        return False
    label = stripped[:-1].strip()
    if not label:
        return False
    return label.isupper() or label.istitle()


def _add_body(document: Document, text: str) -> None:
    """Parse ``text`` line by line and append styled paragraphs."""
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph()
            continue

        if _RULE_RE.match(line):
            _horizontal_rule(document.add_paragraph())
            continue

        if _is_field_label(line):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            run = paragraph.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = NAVY
            continue

        if _NUMBERED_RE.match(line):
            content = _NUMBERED_RE.sub("", line).strip()
            try:
                paragraph = document.add_paragraph(content, style="List Number")
            except KeyError:
                paragraph = document.add_paragraph(line.strip())
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                run.font.size = Pt(11)
            continue

        paragraph = document.add_paragraph(line.strip())
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.size = Pt(11)


def _add_signature_block(document: Document, settings: dict) -> None:
    """Append the closing signature block with a rule and today's date."""
    document.add_paragraph()
    _horizontal_rule(document.add_paragraph())

    prepared = document.add_paragraph()
    run = prepared.add_run(settings.get("prepared_by", ""))
    run.bold = True
    run.font.size = Pt(11)

    if settings.get("title"):
        document.add_paragraph(settings["title"])

    signature = settings.get("signature_block", "")
    for line in signature.splitlines():
        if line.strip():
            document.add_paragraph(line.strip())

    today = datetime.now().strftime("%Y-%m-%d")
    document.add_paragraph(f"Date: {today}")


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def get_output_path(doc_type: str, extension: str) -> str:
    """Return ``outputs/DOCTYPE_YYYYMMDD_HHmmss.ext``, creating the folder."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    safe_type = re.sub(r"[^A-Za-z0-9]+", "", doc_type).upper() or "DOCUMENT"
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ext = extension.lstrip(".")
    return os.path.join(OUTPUT_DIR, f"{safe_type}_{stamp}.{ext}")


def save_as_docx(
    text: str, doc_type: str, output_path: str, settings: dict
) -> str:
    """Render ``text`` into a styled .docx at ``output_path``; return the path."""
    document = Document()
    _style_base_font(document)
    _setup_page(document)
    _build_header(document, doc_type, settings)
    _build_footer(document, settings)
    _build_title_block(document, doc_type, settings)
    _add_body(document, text)
    _add_signature_block(document, settings)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    document.save(output_path)
    return output_path


def _convert_with_libreoffice(docx_path: str, out_dir: str) -> bool:
    """Try converting via a headless LibreOffice install; return success."""
    binary = (
        shutil.which("libreoffice")
        or shutil.which("soffice")
        or _windows_soffice_path()
    )
    if not binary:
        return False
    try:
        subprocess.run(
            [
                binary,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                docx_path,
            ],
            check=True,
            timeout=60,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return True
    except (subprocess.SubprocessError, OSError):
        return False


def _windows_soffice_path() -> str | None:
    """Return a common Windows LibreOffice path if it exists, else None."""
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _convert_with_docx2pdf(docx_path: str, output_path: str) -> bool:
    """Try converting via the docx2pdf library; return success."""
    try:
        from docx2pdf import convert

        convert(docx_path, output_path)
        return os.path.isfile(output_path)
    except Exception:  # noqa: BLE001 - any failure means try the next method.
        return False


def _convert_with_reportlab(
    text: str, doc_type: str, output_path: str, settings: dict
) -> str:
    """Plain-text PDF fallback using reportlab (basic styling only)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    left = MARGIN_MM * mm
    top = height - MARGIN_MM * mm
    bottom = MARGIN_MM * mm
    y = top

    def new_page() -> float:
        pdf.showPage()
        return top

    # Title block.
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawCentredString(width / 2, y, settings.get("contractor_name", ""))
    y -= 8 * mm
    pdf.setFillColorRGB(0x1F / 255, 0x38 / 255, 0x64 / 255)
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawCentredString(width / 2, y, doc_type)
    pdf.setFillColorRGB(0, 0, 0)
    y -= 12 * mm

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if y < bottom:
            y = new_page()
        if not line.strip():
            y -= 5 * mm
            continue
        if _RULE_RE.match(line):
            pdf.line(left, y, width - left, y)
            y -= 5 * mm
            continue
        if _is_field_label(line):
            pdf.setFont("Helvetica-Bold", 11)
            pdf.setFillColorRGB(0x1F / 255, 0x38 / 255, 0x64 / 255)
        else:
            pdf.setFont("Helvetica", 11)
            pdf.setFillColorRGB(0, 0, 0)
        for chunk in _wrap(line.strip(), 95):
            if y < bottom:
                y = new_page()
            pdf.drawString(left, y, chunk)
            y -= 6 * mm

    pdf.setFillColorRGB(0, 0, 0)
    pdf.save()
    print(
        "Note: Using plain text PDF. Install LibreOffice for better formatting."
    )
    return output_path


def _wrap(text: str, width: int) -> list[str]:
    """Greedy word-wrap helper for the reportlab fallback."""
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        if len(current) + len(word) + 1 <= width:
            current = f"{current} {word}".strip()
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def save_as_pdf(
    text: str, doc_type: str, output_path: str, settings: dict
) -> str:
    """Render ``text`` to a PDF at ``output_path`` using the best method.

    Strategy: build a styled .docx then convert with LibreOffice; if that is
    unavailable fall back to docx2pdf, then to a plain reportlab PDF.
    """
    out_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(out_dir, exist_ok=True)

    tmp_docx = ""
    try:
        handle = tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False, dir=out_dir
        )
        tmp_docx = handle.name
        handle.close()
        save_as_docx(text, doc_type, tmp_docx, settings)

        # Method 1 — LibreOffice (outputs <tmpname>.pdf into out_dir).
        if _convert_with_libreoffice(tmp_docx, out_dir):
            produced = os.path.splitext(tmp_docx)[0] + ".pdf"
            if os.path.isfile(produced):
                shutil.move(produced, output_path)
                print("PDF method: LibreOffice")
                return output_path

        # Method 2 — docx2pdf.
        if _convert_with_docx2pdf(tmp_docx, output_path):
            print("PDF method: docx2pdf")
            return output_path

        # Method 3 — reportlab plain-text fallback.
        return _convert_with_reportlab(text, doc_type, output_path, settings)
    finally:
        if tmp_docx and os.path.isfile(tmp_docx):
            try:
                os.unlink(tmp_docx)
            except OSError:
                pass


def save_as_both(
    text: str, doc_type: str, base_path: str, settings: dict
) -> dict:
    """Save matching .docx and .pdf files; return their paths.

    ``base_path`` may be given with or without an extension; the stem is reused
    for both outputs.
    """
    stem = os.path.splitext(base_path)[0]
    docx_path = save_as_docx(text, doc_type, stem + ".docx", settings)
    pdf_path = save_as_pdf(text, doc_type, stem + ".pdf", settings)
    return {"docx": docx_path, "pdf": pdf_path}


if __name__ == "__main__":
    sample_settings = {
        "project_name": "Metro Station A",
        "contractor_name": "Meridian Construction Co.",
        "contract_number": "CN-2026-014",
        "prepared_by": "Jordan Rivera",
        "title": "Project Manager",
        "signature_block": (
            "Meridian Construction Co.\n"
            "Phone: +1 (555) 010-2030\n"
            "Email: jordan.rivera@example.com"
        ),
    }
    sample_text = (
        "REQUEST FOR INFORMATION\n"
        "===\n"
        "RFI NUMBER: RFI-047\n"
        "PROJECT: Metro Station A\n\n"
        "DESCRIPTION:\n"
        "The rebar spacing on structural drawing S-204 conflicts with the "
        "minimum clear spacing required by Specification Section 03 21 00.\n\n"
        "REQUESTED ACTION:\n"
        "1. Confirm the governing spacing for the pile cap.\n"
        "2. Issue a revised detail if required.\n"
        "3. Advise the impact on the pour scheduled for 2026-06-15.\n"
    )

    docx_out = get_output_path("RFI", "docx")
    pdf_out = get_output_path("RFI", "pdf")
    save_as_docx(sample_text, "RFI", docx_out, sample_settings)
    print(f"Saved DOCX: {docx_out}")
    save_as_pdf(sample_text, "RFI", pdf_out, sample_settings)
    print(f"Saved PDF:  {pdf_out}")
